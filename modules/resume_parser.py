"""
resume_parser.py
----------------
Handles extraction of raw text from uploaded resume files.
Supports both PDF (.pdf) and DOCX (.docx) formats.
"""

import io
import re


def extract_text_from_pdf(file) -> str:
    """
    Extract plain text from a PDF file object.

    Args:
        file: File-like object (BytesIO or uploaded file) of a PDF.

    Returns:
        Extracted text as a string, or empty string on failure.
    """
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber is unavailable
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"[ERROR] PyPDF2 extraction failed: {e}")
    except Exception as e:
        print(f"[ERROR] PDF extraction failed: {e}")

    return text.strip()


def extract_text_from_docx(file) -> str:
    """
    Extract plain text from a DOCX file object.

    Args:
        file: File-like object (BytesIO or uploaded file) of a DOCX.

    Returns:
        Extracted text as a string, or empty string on failure.
    """
    text = ""
    try:
        from docx import Document
        document = Document(file)
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)

        # Also extract text from tables inside the document
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed: {e}")

    return text.strip()


def extract_resume_text(uploaded_file) -> str:
    """
    Auto-detect file type and route to appropriate extractor.

    Args:
        uploaded_file: Streamlit UploadedFile object.

    Returns:
        Extracted text string.
    """
    filename = uploaded_file.name.lower()
    file_bytes = io.BytesIO(uploaded_file.read())

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: '{filename}'. Please upload a PDF or DOCX file.")


def extract_contact_info(text: str) -> dict:
    """
    Extract basic contact information from resume text using regex.

    Args:
        text: Raw resume text.

    Returns:
        Dictionary with keys: email, phone, linkedin, github.
    """
    contact = {"email": None, "phone": None, "linkedin": None, "github": None}

    # Email
    email_pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    emails = re.findall(email_pattern, text)
    if emails:
        contact["email"] = emails[0]

    # Phone (supports various formats)
    phone_pattern = r"(\+?\d[\d\s\-().]{8,15}\d)"
    phones = re.findall(phone_pattern, text)
    if phones:
        contact["phone"] = phones[0].strip()

    # LinkedIn
    linkedin_pattern = r"linkedin\.com/in/[\w\-]+"
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    if linkedin:
        contact["linkedin"] = "https://www." + linkedin[0]

    # GitHub
    github_pattern = r"github\.com/[\w\-]+"
    github = re.findall(github_pattern, text, re.IGNORECASE)
    if github:
        contact["github"] = "https://www." + github[0]

    return contact


def estimate_experience_years(text: str) -> int:
    """
    Heuristically estimate years of experience from resume text.

    Args:
        text: Raw resume text.

    Returns:
        Estimated years of experience (0 if not determinable).
    """
    patterns = [
        r"(\d+)\+?\s*years?\s+of\s+experience",
        r"(\d+)\+?\s*years?\s+experience",
        r"experience\s+of\s+(\d+)\+?\s*years?",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return int(matches[0])
    return 0
